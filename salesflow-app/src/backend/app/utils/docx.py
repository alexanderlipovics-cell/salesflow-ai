"""
DOCX Text Extraction Utilities
Extrahiert Text aus Word-Dokumenten für Storybook-Import
"""

from typing import Union
import io


def extract_text_from_docx(content: Union[bytes, str]) -> str:
    """
    Extrahiert Text aus DOCX-Bytes oder Dateipfad
    
    Args:
        content: DOCX als bytes oder Dateipfad
        
    Returns:
        Extrahierter Text
    """
    try:
        return _extract_with_python_docx(content)
    except ImportError:
        pass
    
    try:
        return _extract_with_docx2txt(content)
    except ImportError:
        pass
    
    raise ImportError(
        "Kein DOCX-Parser verfügbar. Installiere eines der folgenden Pakete: "
        "python-docx (pip install python-docx), "
        "docx2txt (pip install docx2txt)"
    )


def _extract_with_python_docx(content: Union[bytes, str]) -> str:
    """Extrahiert Text mit python-docx"""
    from docx import Document
    
    if isinstance(content, bytes):
        doc_file = io.BytesIO(content)
    else:
        doc_file = content
    
    doc = Document(doc_file)
    
    text_parts = []
    
    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            # Check for headings
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                text_parts.append(f"\n{'#' * int(level) if level.isdigit() else '#'} {para.text}\n")
            else:
                text_parts.append(para.text)
    
    # Tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))
        if table_text:
            text_parts.append("\n[Tabelle]\n" + "\n".join(table_text) + "\n")
    
    return "\n".join(text_parts)


def _extract_with_docx2txt(content: Union[bytes, str]) -> str:
    """Extrahiert Text mit docx2txt"""
    import docx2txt
    
    if isinstance(content, bytes):
        doc_file = io.BytesIO(content)
        text = docx2txt.process(doc_file)
    else:
        text = docx2txt.process(content)
    
    return text


def get_docx_metadata(content: Union[bytes, str]) -> dict:
    """
    Extrahiert Metadaten aus DOCX
    
    Returns:
        Dict mit title, author, created, modified, etc.
    """
    try:
        from docx import Document
        
        if isinstance(content, bytes):
            doc_file = io.BytesIO(content)
        else:
            doc_file = content
        
        doc = Document(doc_file)
        props = doc.core_properties
        
        return {
            "title": props.title,
            "author": props.author,
            "subject": props.subject,
            "keywords": props.keywords,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "last_modified_by": props.last_modified_by,
            "revision": props.revision,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
    except ImportError:
        return {"error": "python-docx not available"}
    except Exception as e:
        return {"error": str(e)}


def extract_text_with_structure(content: Union[bytes, str]) -> dict:
    """
    Extrahiert Text mit Struktur-Informationen
    
    Returns:
        Dict mit sections, headings, paragraphs, tables
    """
    try:
        from docx import Document
        
        if isinstance(content, bytes):
            doc_file = io.BytesIO(content)
        else:
            doc_file = content
        
        doc = Document(doc_file)
        
        result = {
            "sections": [],
            "headings": [],
            "paragraphs": [],
            "tables": [],
            "full_text": "",
        }
        
        current_section = {"heading": None, "content": []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check for headings
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_section["heading"] or current_section["content"]:
                    result["sections"].append(current_section)
                
                level = para.style.name.replace('Heading ', '')
                result["headings"].append({
                    "level": int(level) if level.isdigit() else 1,
                    "text": text,
                })
                
                current_section = {"heading": text, "level": level, "content": []}
            else:
                result["paragraphs"].append(text)
                current_section["content"].append(text)
        
        # Add last section
        if current_section["heading"] or current_section["content"]:
            result["sections"].append(current_section)
        
        # Tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result["tables"].append(table_data)
        
        # Full text
        result["full_text"] = "\n\n".join(result["paragraphs"])
        
        return result
        
    except ImportError:
        raise ImportError("python-docx nicht verfügbar. Installiere: pip install python-docx")

